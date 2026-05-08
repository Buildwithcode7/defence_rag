"""
dev_server.py - INICAI DefenceRAG v5.0 — Semantic + Keyword Hybrid Search
Architecture:
  1. Full page extraction (pdfplumber tables + pymupdf text)
  2. FAISS semantic search (meaning-based, finds related content even if words differ)
  3. BM25 keyword search (exact term matching)
  4. RRF fusion (combines both for best results)
  5. LLM answers from full retrieved pages

Install:
  pip install fastapi uvicorn[standard] pymupdf pdfplumber
              sentence-transformers faiss-cpu python-multipart numpy rank-bm25

Run: python dev_server.py
"""
import os, uuid, time, base64, json, re, pickle, logging
from pathlib import Path
from typing import List, Optional, Dict, Tuple
from collections import defaultdict
import os

# _env = Path(__file__).parent / ".env"
# if _env.exists():
#     for _l in _env.read_text().splitlines():
#         _l = _l.strip()
#         if _l and not _l.startswith("#") and "=" in _l:
#             _k, _v = _l.split("=", 1)
#             os.environ.setdefault(_k.strip(), _v.strip().strip('"').strip("'"))

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")
logger = logging.getLogger(__name__)
for _d in ["data/raw", "data/indices", "data/audit"]:
    Path(_d).mkdir(parents=True, exist_ok=True)

# =============================================================================
# PDF / DOCUMENT EXTRACTION  (full text + tables, nothing lost)
# =============================================================================

def _clean(t: str) -> str:
    t = re.sub(r"\n{3,}", "\n\n", t)
    return re.sub(r"[ \t]{2,}", " ", t).strip()

def _rows_to_md(rows) -> str:
    if not rows: return ""
    out = []
    for ri, row in enumerate(rows):
        cells = [str(c).strip() if c else "" for c in row]
        non_empty = [c for c in cells if c]
        if not non_empty: continue
        if ri == 0:
            out.append("| " + " | ".join(cells) + " |")
            out.append("|" + "|".join(["---"] * len(cells)) + "|")
        else:
            out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


# =============================================================================
# LOCATION EXTRACTOR — finds paragraph number and exact snippet for any answer
# =============================================================================

def split_paragraphs(text: str) -> List[str]:
    """Split page text into numbered paragraphs."""
    paras = re.split(r'\n{2,}', text)
    return [p.strip() for p in paras if p.strip() and len(p.strip()) > 10]


def find_answer_location(answer_text: str, page_text: str) -> dict:
    """
    Given an answer string and the full page text, find:
    - paragraph_number: which paragraph contains the answer
    - paragraph_text:   the actual paragraph text
    - char_start:       character offset in page
    - snippet:          40-word context window around the match
    """
    if not answer_text or not page_text:
        return {}

    paras = split_paragraphs(page_text)

    # Strategy 1: find longest common substring (>15 chars) between answer and page
    # Slide a window over answer words and search in page
    answer_words = answer_text.lower().split()
    best_para_idx  = -1
    best_para_text = ""
    best_char_start = -1
    best_snippet    = ""
    best_score      = 0

    # Try windows of 8, 5, 3 words
    for window in [8, 5, 3]:
        if len(answer_words) < window:
            continue
        for start_w in range(0, min(len(answer_words) - window + 1, 30)):
            phrase = " ".join(answer_words[start_w: start_w + window])
            if len(phrase) < 10:
                continue
            # Search in page text (case-insensitive)
            page_lower = page_text.lower()
            idx = page_lower.find(phrase)
            if idx == -1:
                continue
            # Found — which paragraph?
            char_count = 0
            for pi, para in enumerate(paras):
                if page_text.lower().find(phrase) >= char_count and \
                   page_text.lower().find(phrase) < char_count + len(para) + 4:
                    if window > best_score:
                        best_score      = window
                        best_para_idx   = pi
                        best_para_text  = para
                        best_char_start = idx
                        # Extract snippet: 40 words around the match
                        words_before = page_text[:idx].split()[-10:]
                        words_after  = page_text[idx + len(phrase):].split()[:20]
                        match_words  = phrase.split()
                        best_snippet = " ".join(words_before + match_words + words_after)
                    break
                char_count += len(para) + 2
        if best_score >= 5:
            break

    # Strategy 2: keyword overlap scoring per paragraph
    if best_para_idx == -1:
        answer_kws = set(re.findall(r"[a-zA-Z0-9]{4,}", answer_text.lower()))
        best_overlap = 0
        for pi, para in enumerate(paras):
            para_kws = set(re.findall(r"[a-zA-Z0-9]{4,}", para.lower()))
            overlap  = len(answer_kws & para_kws)
            if overlap > best_overlap:
                best_overlap   = overlap
                best_para_idx  = pi
                best_para_text = para
                # Snippet: first 150 chars of paragraph
                best_snippet   = para[:200]

    if best_para_idx == -1:
        return {}

    return {
        "paragraph_number": best_para_idx + 1,          # 1-indexed
        "paragraph_text":   best_para_text[:300],        # first 300 chars
        "char_start":       max(best_char_start, 0),
        "snippet":          best_snippet[:300],
        "total_paragraphs": len(paras),
    }


def enrich_citations_with_locations(
    citations: List[dict],
    pages: List[dict],
    answer: str,
) -> List[dict]:
    """
    Post-process citations to add paragraph and location data.
    Matches the LLM answer text back to the source pages.
    """
    # Split answer into sentences to match per-source
    answer_sentences = re.split(r"(?<=[.!?])\s+", answer)

    for i, (cit, page) in enumerate(zip(citations, pages)):
        page_text = page.get("text", "")
        paras     = split_paragraphs(page_text)

        # Try to find which part of the answer came from this page
        # Use the text_snippet as the search target
        search_text = cit.get("text_snippet", page_text[:200])
        location    = find_answer_location(search_text, page_text)

        cit["paragraph_number"]  = location.get("paragraph_number", "N/A")
        cit["paragraph_text"]    = location.get("paragraph_text",   "")
        cit["total_paragraphs"]  = location.get("total_paragraphs", len(paras))
        cit["exact_snippet"]     = location.get("snippet",          page_text[:200])
        cit["char_offset"]       = location.get("char_start",       0)

        # Build human-readable location string
        pg  = cit.get("page", "?")
        par = cit.get("paragraph_number", "?")
        cit["location_label"] = f"Page {pg}, Paragraph {par}"

    return citations


def extract_pdf(path: str) -> Dict[int, str]:
    import fitz
    pages: Dict[int, str] = {}
    plumber_tables: Dict[int, str] = {}

    # Step 1: extract tables with pdfplumber (best table parser)
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                tbls = page.extract_tables({
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 5,
                })
                if tbls:
                    plumber_tables[i] = "\n\n".join(
                        _rows_to_md(t) for t in tbls if t
                    )
    except Exception as e:
        logger.warning("pdfplumber failed: %s", e)

    # Step 2: extract text blocks with pymupdf
    doc = fitz.open(path)
    for pnum, page in enumerate(doc, 1):
        # Get all text blocks sorted top-to-bottom, left-to-right
        blocks = page.get_text("blocks")
        text_parts = [
            b[4].strip()
            for b in sorted(blocks, key=lambda b: (round(b[1] / 10) * 10, b[0]))
            if b[4].strip() and len(b[4].strip()) > 2
        ]
        raw = "\n".join(text_parts)

        # Try fitz table extraction too
        fitz_tables = ""
        try:
            tabs = page.find_tables()
            if tabs and tabs.tables:
                fitz_tables = "\n\n".join(
                    _rows_to_md(t.extract()) for t in tabs.tables if t.extract()
                )
        except Exception:
            pass

        # Prefer pdfplumber tables, fallback to fitz
        table_text = plumber_tables.get(pnum, fitz_tables)

        if table_text:
            # Clean raw text: remove lines that are just separators (table artifacts)
            clean_raw = "\n".join(
                ln for ln in raw.splitlines()
                if ln.strip() and not re.match(r"^[\s|—\-_=]+$", ln.strip())
            )
            combined = clean_raw + "\n\n" + table_text
        else:
            combined = raw

        combined = _clean(combined)
        if combined and len(combined) > 20:
            pages[pnum] = combined

    doc.close()
    logger.info("PDF extracted: %d pages with content", len(pages))
    return pages


def extract_doc(path: str) -> Dict[int, str]:
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc  = Document(path)
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for t in doc.tables:
                rows = [[c.text.strip() for c in r.cells] for r in t.rows]
                text += "\n\n" + _rows_to_md(rows)
            return {1: _clean(text)}
        except ImportError:
            raise ImportError("pip install python-docx")
    elif suffix == ".txt":
        return {1: _clean(Path(path).read_text(encoding="utf-8", errors="ignore"))}
    raise ValueError(f"Unsupported file type: {suffix}")


# =============================================================================
# DOCUMENT STORE  — stores full text of every page
# =============================================================================

class DocStore:
    PATH = "data/indices/docstore_v5.pkl"

    def __init__(self):
        self._docs:  Dict[str, dict] = {}
        self._pages: List[dict]      = []
        if Path(self.PATH).exists():
            try:
                with open(self.PATH, "rb") as f:
                    d = pickle.load(f)
                self._docs  = d.get("docs", {})
                self._pages = d.get("pages", [])
                logger.info("DocStore loaded: %d docs, %d pages", len(self._docs), len(self._pages))
            except Exception as e:
                logger.warning("DocStore load failed: %s", e)

    def _save(self):
        with open(self.PATH, "wb") as f:
            pickle.dump({"docs": self._docs, "pages": self._pages}, f)

    @property
    def total_docs(self):  return len(self._docs)
    @property
    def total_pages(self): return len(self._pages)

    def add(self, doc_id: str, filename: str, page_texts: Dict[int, str], meta: dict) -> int:
        full = "\n\n".join(f"[PAGE {n}]\n{t}" for n, t in sorted(page_texts.items()))
        self._docs[doc_id] = {
            "filename":   filename,
            "pages":      page_texts,
            "full_text":  full,
            "meta":       meta,
            "page_count": len(page_texts),
        }
        for pnum, ptext in page_texts.items():
            self._pages.append({
                "page_id":  f"{doc_id}__p{pnum}",
                "doc_id":   doc_id,
                "filename": filename,
                "page":     pnum,
                "text":     ptext,
            })
        self._save()
        return len(page_texts)

    def pages(self):     return self._pages
    def list_docs(self): return [
        {"doc_id": k, "filename": v["filename"], "page_count": v["page_count"], "meta": v["meta"]}
        for k, v in self._docs.items()
    ]


# =============================================================================
# EMBEDDER  — sentence-transformers for semantic (meaning-based) search
# =============================================================================

class Embedder:
    """
    Uses multi-qa-MiniLM-L6-cos-v1 — trained specifically for Q&A retrieval.
    Finds passages that MEAN the same thing as the query, even if different words.
    Example: query "credit hours" finds "units of study" without exact word match.
    """
    DIM   = 384
    MODEL = "multi-qa-MiniLM-L6-cos-v1"

    def __init__(self):
        self._model = None
        try:
            from sentence_transformers import SentenceTransformer
            logger.info("Loading semantic embedding model: %s ...", self.MODEL)
            self._model = SentenceTransformer(self.MODEL)
            try:    self.DIM = self._model.get_embedding_dimension()
            except: self.DIM = self._model.get_sentence_embedding_dimension()
            logger.info("Semantic model ready. dim=%d", self.DIM)
        except Exception as e:
            logger.warning("Embedding model unavailable (%s) — semantic search disabled", e)

    def embed(self, text: str) -> Optional[List[float]]:
        if self._model is None: return None
        return self._model.encode(
            [text], show_progress_bar=False,
            convert_to_numpy=True, normalize_embeddings=True
        )[0].tolist()

    def embed_batch(self, texts: List[str]) -> Optional[List[List[float]]]:
        if self._model is None: return None
        return self._model.encode(
            texts, batch_size=32, show_progress_bar=False,
            convert_to_numpy=True, normalize_embeddings=True
        ).tolist()

    @property
    def available(self): return self._model is not None


# =============================================================================
# FAISS SEMANTIC INDEX
# =============================================================================

class SemanticIndex:
    """
    FAISS index for semantic (meaning-based) search.
    Finds documents semantically related to the query — even with different words.
    """
    F = "data/indices/faiss_v5.index"
    M = "data/indices/faiss_v5_meta.pkl"

    def __init__(self, dim: int = 384):
        self.dim   = dim
        self._idx  = None
        self._meta: List[dict] = []
        if Path(self.F).exists() and Path(self.M).exists():
            try:
                import faiss
                self._idx = faiss.read_index(self.F)
                with open(self.M, "rb") as f:
                    self._meta = pickle.load(f)
                logger.info("FAISS semantic index loaded: %d vectors", len(self._meta))
            except Exception as e:
                logger.warning("FAISS load failed: %s", e)
        else:
            logger.info("No FAISS index — will build on first ingest")

    def add(self, embeddings: List[List[float]], metadatas: List[dict]):
        import faiss, numpy as np
        v = np.array(embeddings, dtype="float32")
        if self._idx is None:
            # IndexFlatIP = inner product (cosine similarity when vectors are normalised)
            self._idx = faiss.IndexFlatIP(self.dim)
        self._idx.add(v)
        self._meta.extend(metadatas)
        faiss.write_index(self._idx, self.F)
        with open(self.M, "wb") as f:
            pickle.dump(self._meta, f)
        logger.info("FAISS: %d total vectors", len(self._meta))

    def search(self, q_emb: List[float], top_k: int = 10) -> List[Tuple[dict, float]]:
        import numpy as np
        if self._idx is None or self._idx.ntotal == 0: return []
        k = min(top_k, self._idx.ntotal)
        q = np.array([q_emb], dtype="float32")
        scores, idxs = self._idx.search(q, k)
        return [
            (self._meta[i], float(s))
            for s, i in zip(scores[0], idxs[0])
            if 0 <= i < len(self._meta) and s > 0
        ]

    def rebuild(self, pages: List[dict], embedder: Embedder):
        """Rebuild entire FAISS index from page list."""
        if not embedder.available or not pages:
            return
        logger.info("Rebuilding FAISS index for %d pages...", len(pages))
        import faiss
        self._idx  = faiss.IndexFlatIP(self.dim)
        self._meta = []
        texts      = [p["text"] for p in pages]
        embeddings = embedder.embed_batch(texts)
        if embeddings:
            import numpy as np
            v = np.array(embeddings, dtype="float32")
            self._idx.add(v)
            self._meta = [dict(p) for p in pages]
            faiss.write_index(self._idx, self.F)
            with open(self.M, "wb") as f:
                pickle.dump(self._meta, f)
            logger.info("FAISS rebuilt: %d vectors", len(self._meta))


# =============================================================================
# BM25 KEYWORD INDEX
# =============================================================================

class BM25Index:
    """
    BM25 keyword search — finds exact term matches.
    Best for: course codes (BAI401), rule numbers, specific names.
    """
    P = "data/indices/bm25_v5.pkl"

    def __init__(self):
        self._bm25 = None
        self._meta: List[dict] = []
        if Path(self.P).exists():
            try:
                with open(self.P, "rb") as f:
                    saved = pickle.load(f)
                self._bm25 = saved["bm25"]
                self._meta = saved["meta"]
                logger.info("BM25 loaded: %d pages", len(self._meta))
            except Exception as e:
                logger.warning("BM25 load failed: %s", e)

    @staticmethod
    def _tok(text: str) -> List[str]:
        # Tokenize: keep alphanumeric + hyphens (preserves course codes like BAI401)
        return [t for t in re.findall(r"[a-z0-9]+(?:-[a-z0-9]+)*", text.lower()) if len(t) > 1]

    def rebuild(self, pages: List[dict]):
        try:
            from rank_bm25 import BM25Okapi
        except ImportError:
            logger.warning("rank-bm25 not installed — pip install rank-bm25"); return
        if not pages: return
        self._bm25 = BM25Okapi([self._tok(p["text"]) for p in pages])
        self._meta = [dict(p) for p in pages]
        with open(self.P, "wb") as f:
            pickle.dump({"bm25": self._bm25, "meta": self._meta}, f)
        logger.info("BM25 rebuilt: %d pages", len(self._meta))

    def search(self, query: str, top_k: int = 10) -> List[Tuple[dict, float]]:
        if not self._bm25 or not self._meta: return []
        toks = self._tok(query)
        if not toks: return []
        scores = self._bm25.get_scores(toks)
        top_i  = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:top_k]
        return [(self._meta[i], float(scores[i])) for i in top_i if scores[i] > 0]


# =============================================================================
# HYBRID RETRIEVER  — RRF fusion of semantic + keyword results
# =============================================================================

class HybridRetriever:
    """
    Combines:
    - Semantic search (FAISS): finds related meaning
    - Keyword search (BM25):   finds exact terms
    - RRF fusion:              ranks by both signals together

    This means a query like "subjects in fourth semester" will find pages
    that contain SEMESTER IV tables even if they use different words.
    """

    def __init__(self, store: DocStore, embedder: Embedder):
        self.store    = store
        self.embedder = embedder
        self.faiss    = SemanticIndex(dim=embedder.DIM)
        self.bm25     = BM25Index()

    def rebuild(self):
        pages = self.store.pages()
        if not pages:
            logger.info("No pages to index yet")
            return
        logger.info("Building search indices for %d pages...", len(pages))
        self.faiss.rebuild(pages, self.embedder)
        self.bm25.rebuild(pages)
        logger.info("Hybrid index ready")

    def search(self, query: str, top_k: int = 8, is_broad: bool = False) -> List[dict]:
        k = min(top_k * 2 if is_broad else top_k + 4, self.store.total_pages)

        # --- Semantic search (meaning-based) ---
        sem_hits: List[Tuple[dict, float]] = []
        if self.embedder.available:
            # Search with original query
            q_emb = self.embedder.embed(query)
            if q_emb:
                sem_hits = self.faiss.search(q_emb, top_k=k)

            # Also search with keyword-stripped version for better recall
            q_stripped = re.sub(
                r'^(what|who|when|where|how|is|are|does|do|list|give|find|tell me|show)\s+',
                '', query.lower().rstrip("?")
            ).strip()
            if q_stripped and q_stripped != query.lower():
                q_emb2 = self.embedder.embed(q_stripped)
                if q_emb2:
                    extra = self.faiss.search(q_emb2, top_k=k)
                    # Merge deduplicating by page_id
                    seen = {m.get("page_id") for m, _ in sem_hits}
                    for m, s in extra:
                        if m.get("page_id") not in seen:
                            sem_hits.append((m, s))
                            seen.add(m.get("page_id"))

        # --- Keyword search (exact terms) ---
        kw_hits: List[Tuple[dict, float]] = self.bm25.search(query, top_k=k)

        # --- RRF Fusion ---
        scores:   Dict[str, float] = defaultdict(float)
        meta_map: Dict[str, dict]  = {}
        K = 60

        for rank, (m, _) in enumerate(sem_hits, 1):
            pid = m.get("page_id", f"s{rank}")
            scores[pid]   += 0.6 / (K + rank)   # semantic weight: 60%
            meta_map[pid]  = m

        for rank, (m, _) in enumerate(kw_hits, 1):
            pid = m.get("page_id", f"k{rank}")
            scores[pid]   += 0.4 / (K + rank)   # keyword weight: 40%
            meta_map[pid]  = m

        ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)

        # Fallback: if nothing matched, return top pages from store
        if not ranked:
            logger.warning("No results from hybrid search — returning top pages as fallback")
            return [dict(p, score=0.0) for p in self.store.pages()[:top_k]]

        result = []
        for pid, score in ranked[:top_k * 2 if is_broad else top_k]:
            page = dict(meta_map[pid])
            page["score"]          = round(score, 6)
            page["semantic_score"] = round(
                next((s for m, s in sem_hits if m.get("page_id") == pid), 0.0), 4
            )
            page["keyword_score"]  = round(
                next((s for m, s in kw_hits  if m.get("page_id") == pid), 0.0), 4
            )
            result.append(page)

        logger.info(
            "Hybrid search: %d semantic hits + %d keyword hits → %d fused results",
            len(sem_hits), len(kw_hits), len(result)
        )
        return result


# =============================================================================
# BROAD QUERY DETECTOR
# =============================================================================

_BROAD_RE = re.compile(
    r"\b(all|every|list|complete|full|entire|semester|subjects?|courses?|"
    r"syllabus|curriculum|overview|summary|what are|how many|total|which|"
    r"chapters?|sections?|topics?|units?|modules?|policies|rules|regulations)\b",
    re.IGNORECASE,
)

def is_broad_query(q: str) -> bool:
    return bool(_BROAD_RE.search(q))


# =============================================================================
# LLM
# =============================================================================

SYSTEM_PROMPT = """You are an expert document analyst with perfect reading comprehension.

RULES — follow exactly:
1. Answer ONLY from the document content provided. No outside knowledge.
2. Give the DIRECT answer first — no preamble, no "based on the sources".
3. LIST questions: extract EVERY single item — never stop early, never skip.
4. TABLE data: reproduce as a clean markdown table with ALL rows and columns.
5. Copy course codes, rule numbers, and names EXACTLY as they appear.
6. If info is absent, say: "Not found in the document."
7. After your answer, add a LOCATION section in this exact format:

📍 SOURCE LOCATIONS:
- [SOURCE-N] → Page <number>, Paragraph <number>: "<exact quote of 10-15 words from that page>"

8. For each SOURCE you used, include one location entry.
9. Be complete — if there are 15 courses, list all 15.
10. The exact quote in the location must be copied word-for-word from the source text."""


def call_llm(prompt: str) -> Optional[str]:
    key  = os.getenv("OPENAI_API_KEY", "").strip()
    base = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1").rstrip("/")
    mdl  = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

    if not key:
        logger.warning("OPENAI_API_KEY not set")
        return None

    logger.info("LLM → %s | %s", base, mdl)
    try:
        import requests as rq
        hdrs = {
            "Authorization": f"Bearer {key}",
            "Content-Type":  "application/json",
        }
        if "openrouter" in base:
            hdrs["HTTP-Referer"] = "http://localhost:8000"
            hdrs["X-Title"]      = "INICAI DefenceRAG"

        r = rq.post(
            f"{base}/chat/completions",
            headers=hdrs,
            timeout=None,
            json={
                "model":       mdl,
                "messages":    [
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user",   "content": prompt},
                ],
                "max_tokens":  3000,
                "temperature": 0.0,
            },
        )
        if r.status_code != 200:
            logger.error("LLM error %d: %s", r.status_code, r.text[:400])
            return None
        ans = r.json()["choices"][0]["message"]["content"].strip()
        logger.info("LLM response: %d chars — %s...", len(ans), ans[:80])
        return ans
    except Exception as e:
        logger.error("LLM call failed: %s", e, exc_info=True)
        return None



def _parse_llm_locations_into_citations(answer: str, citations: List[dict]) -> None:
    """
    Parse the 📍 SOURCE LOCATIONS block that the LLM writes into the answer,
    and backfill paragraph_number + exact_snippet into the matching citation.

    Expected format from LLM:
      📍 SOURCE LOCATIONS:
      - [SOURCE-1] → Page 12, Paragraph 3: "exact quote here"
      - [SOURCE-2] → Page 7, Paragraph 1: "another exact quote"
    """
    # Find the locations block
    loc_match = re.search(r"📍 SOURCE LOCATIONS:?(.*?)(?:\Z|\n\n(?:[^\-]))", answer, re.DOTALL | re.IGNORECASE)
    if not loc_match:
        loc_match = re.search(r"SOURCE LOCATIONS:?(.*?)$", answer, re.DOTALL | re.IGNORECASE)
    if not loc_match:
        return

    block = loc_match.group(1)
    # Parse each line: - [SOURCE-N] → Page X, Paragraph Y: "quote"
    pattern = re.compile(
        r'\[SOURCE-?(\d+)\].*?Page\s+(\d+).*?Paragraph\s+(\d+).*?[:=-]\s*.{0,5}(.{8,80})',
        re.IGNORECASE,
    )
    for m in pattern.finditer(block):
        src_num  = int(m.group(1)) - 1   # 0-indexed
        page_num = m.group(2)
        para_num = m.group(3)
        quote    = m.group(4).strip().strip('"').strip("\u201c\u201d")

        if 0 <= src_num < len(citations):
            cit = citations[src_num]
            cit["paragraph_number"] = int(para_num)
            cit["exact_snippet"]    = quote
            cit["location_label"]   = f"Page {page_num}, Paragraph {para_num}"
            logger.info(
                "Location parsed: SOURCE-%d → Page %s ¶%s: '%s'",
                src_num + 1, page_num, para_num, quote[:40],
            )

# =============================================================================
# RAG PIPELINE
# =============================================================================

class RAGPipeline:
    def __init__(self):
        self.store     = DocStore()
        self.embedder  = Embedder()
        self.retriever = HybridRetriever(self.store, self.embedder)
        # If we already have docs, rebuild search indices
        if self.store.total_pages > 0:
            logger.info("Rebuilding search indices from existing data...")
            self.retriever.rebuild()
        logger.info(
            "RAGPipeline v5 ready. Docs: %d | Pages: %d | Semantic: %s",
            self.store.total_docs,
            self.store.total_pages,
            "ON" if self.embedder.available else "OFF (pip install sentence-transformers)",
        )

    # ------------------------------------------------------------------
    def ingest(self, file_path: str, meta: dict = None) -> dict:
        doc_id   = str(uuid.uuid4())
        meta     = meta or {}
        filename = Path(file_path).name

        logger.info("Extracting '%s'...", filename)
        page_texts = extract_doc(file_path)
        if not page_texts:
            raise ValueError("No text could be extracted from this file.")

        total_chars = sum(len(t) for t in page_texts.values())
        n = self.store.add(doc_id, filename, page_texts, meta)

        # Rebuild BOTH search indices after every ingest
        self.retriever.rebuild()

        logger.info("Ingested '%s': %d pages, %d chars", filename, n, total_chars)
        return {
            "doc_id":      doc_id,
            "pages_indexed": n,
            "total_chars": total_chars,
            "total_docs":  self.store.total_docs,
            "total_pages": self.store.total_pages,
        }

    # ------------------------------------------------------------------
    def query(self, question: str, top_k: int = 8) -> dict:
        aid  = str(uuid.uuid4())

        if self.store.total_docs == 0:
            return self._empty("No documents uploaded yet. Please upload a PDF first.", aid)

        broad = is_broad_query(question)
        pages = self.retriever.search(question, top_k=top_k, is_broad=broad)

        if not pages:
            return self._empty("No relevant pages found.", aid)

        # Build context — full page text with paragraph numbers, no truncation
        context_parts = []
        citations     = []
        for i, page in enumerate(pages, 1):
            label = f"SOURCE-{i}"
            text  = page["text"]
            sem_s = page.get("semantic_score", 0)
            kw_s  = page.get("keyword_score",  0)
            paras = split_paragraphs(text)

            # Number each paragraph in the context so LLM can cite them
            numbered_text = ""
            for pi, para in enumerate(paras, 1):
                numbered_text += f"[¶{pi}] {para}\n\n"

            context_parts.append(
                f"[{label}] {page['filename']} — Page {page['page']} "
                f"| {len(paras)} paragraphs | "
                f"(semantic:{sem_s:.3f} keyword:{kw_s:.2f})\n"
                f"{'─'*55}\n{numbered_text.strip()}"
            )
            citations.append({
                "label":            label,
                "doc_id":           page["doc_id"],
                "section":          page["filename"],
                "page":             page["page"],
                "score":            round(page.get("score", 0), 6),
                "semantic_score":   sem_s,
                "keyword_score":    kw_s,
                "text_snippet":     text[:400],
                "total_paragraphs": len(paras),
                "paragraph_number": "N/A",   # filled after LLM call
                "exact_snippet":    "",       # filled after LLM call
                "location_label":   f"Page {page['page']}",
            })

        context     = ("\n\n" + "═" * 60 + "\n\n").join(context_parts)
        docs_list   = ", ".join(f"'{d['filename']}'" for d in self.store.list_docs())
        total_chars = sum(len(p["text"]) for p in pages)

        logger.info(
            "Query '%s' → %d pages, %d chars, broad=%s, semantic=%s",
            question[:60], len(pages), total_chars, broad,
            self.embedder.available,
        )

        prompt = (
            f"Documents: {docs_list}\n"
            f"Pages retrieved: {len(pages)} (broad query: {broad})\n"
            f"Note: Each paragraph is numbered [¶N] so you can cite exact locations.\n\n"
            f"{'═'*60}\n"
            f"{context}\n"
            f"{'═'*60}\n\n"
            f"QUESTION: {question}\n\n"
            f"Instructions:\n"
            f"1. Answer directly and completely.\n"
            f"2. After the answer, add a 📍 SOURCE LOCATIONS section listing:\n"
            f"   - Which SOURCE (SOURCE-1, SOURCE-2 etc) you used\n"
            f"   - The Page number\n"
            f"   - The Paragraph number (¶N from the numbered text above)\n"
            f"   - An exact 10-15 word quote from that paragraph\n"
        )

        answer = call_llm(prompt)

        # Enrich citations with paragraph locations matched from the answer
        if answer:
            citations = enrich_citations_with_locations(citations, pages, answer)
            # Also parse LLM-reported locations from the answer text
            _parse_llm_locations_into_citations(answer, citations)

        if not answer:
            key_ok = bool(os.getenv("OPENAI_API_KEY", "").strip())
            if not key_ok:
                answer = "⚠️ No API key configured. Add OPENAI_API_KEY to your .env file."
            else:
                answer = (
                    "⚠️ LLM call failed. Check your API key.\n\n"
                    "**Retrieved pages:**\n\n" +
                    "\n\n---\n\n".join(
                        f"**Page {p['page']} — {p['filename']}**\n{p['text'][:800]}"
                        for p in pages
                    )
                )

        conf  = min(sum(p.get("score", 0) for p in pages) / max(len(pages), 1) * 500, 1.0)
        level = "HIGH" if conf > 0.6 else ("MEDIUM" if conf > 0.3 else "LOW")

        return {
            "answer":           answer,
            "annotated_answer": answer,
            "citations":        citations,
            "compliance_status": "REVIEW REQUIRED",
            "compliance_gaps":  [],
            "applicable_rules": [],
            "confidence_score": round(conf, 3),
            "confidence_level": level,
            "cot_applied":      False,
            "audit_id":         aid,
            "session_id":       None,
            "pages_searched":   len(pages),
            "context_chars":    total_chars,
            "semantic_enabled": self.embedder.available,
        }

    def _empty(self, msg: str, aid: str) -> dict:
        return {
            "answer": msg, "annotated_answer": msg, "citations": [],
            "compliance_status": "INSUFFICIENT_BASIS", "compliance_gaps": [],
            "applicable_rules": [], "confidence_score": 0.0, "confidence_level": "LOW",
            "cot_applied": False, "audit_id": aid, "session_id": None,
        }


# =============================================================================
# FASTAPI
# =============================================================================

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

app = FastAPI(title="INICAI DefenceRAG API", version="5.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True,
                   allow_methods=["*"], allow_headers=["*"])

@app.middleware("http")
async def timing_mw(req: Request, call_next):
    t = time.time()
    r = await call_next(req)
    r.headers["X-Process-Time"] = f"{(time.time()-t)*1000:.1f}ms"
    return r

_pipeline: Optional[RAGPipeline] = None

def get_pipeline() -> RAGPipeline:
    global _pipeline
    if _pipeline is None:
        _pipeline = RAGPipeline()
    return _pipeline

@app.on_event("startup")
def startup(): get_pipeline()

USERS = {
    "analyst":  ("analyst123",  "analyst"),
    "admin":    ("admin123",    "admin"),
    "auditor":  ("auditor123",  "auditor"),
}

class LoginReq(BaseModel): username: str; password: str

@app.post("/api/v1/auth/login")
def login(body: LoginReq):
    u = USERS.get(body.username)
    if not u or u[0] != body.password: raise HTTPException(401, "Invalid credentials")
    token = base64.b64encode(json.dumps({"sub": body.username, "role": u[1]}).encode()).decode()
    return {"access_token": token, "token_type": "bearer", "role": u[1], "expires_in_minutes": 480}

@app.get("/api/v1/health")
def health():
    p   = get_pipeline()
    key = bool(os.getenv("OPENAI_API_KEY", "").strip())
    return {
        "status":                "healthy",
        "llm_model":             os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        "llm_service_ready":     key,
        "semantic_search":       p.embedder.available,
        "embedding_model":       p.embedder.MODEL if p.embedder.available else "not loaded",
        "total_documents":       p.store.total_docs,
        "total_pages":           p.store.total_pages,
        "faiss_vectors":         len(p.retriever.faiss._meta),
        "bm25_pages":            len(p.retriever.bm25._meta),
        "documents":             p.store.list_docs(),
        "version":               "5.0.0",
    }

class QueryReq(BaseModel):
    question:   str
    filters:    Optional[dict] = None
    session_id: Optional[str]  = None
    top_k:      int            = 8

@app.post("/api/v1/query")
def query_ep(body: QueryReq):
    logger.info("QUERY: %s", body.question[:100])
    try:
        r = get_pipeline().query(body.question, top_k=body.top_k)
        r["session_id"] = body.session_id
        return r
    except Exception as e:
        logger.error("Query error: %s", e, exc_info=True)
        raise HTTPException(500, str(e))

@app.post("/api/v1/ingest")
async def ingest_ep(
    file:                 UploadFile = File(...),
    doc_type:             str        = Form("auto"),
    classification_level: str        = Form("UNCLASSIFIED"),
    issuing_authority:    str        = Form(""),
    effective_date:       str        = Form(""),
):
    logger.info("INGEST: %s", file.filename)
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt"}:
        raise HTTPException(400, f"Unsupported: {suffix}. Use PDF, DOCX, or TXT.")
    content = await file.read()
    if not content:                    raise HTTPException(400, "File is empty.")
    if len(content) > 100*1024*1024:   raise HTTPException(413, "Too large (max 100 MB).")

    job_id = str(uuid.uuid4())
    path   = Path("data/raw") / f"{job_id}_{file.filename}"
    path.write_bytes(content)

    try:
        r = get_pipeline().ingest(str(path), {
            "doc_type":             doc_type,
            "classification_level": classification_level,
            "issuing_authority":    issuing_authority,
            "effective_date":       effective_date,
        })
        return {
            "job_id":        job_id,
            "filename":      file.filename,
            "status":        "complete",
            "message":       (
                f"✅ '{file.filename}' fully analysed! "
                f"{r['pages_indexed']} pages, {r['total_chars']:,} chars. "
                f"Semantic + keyword search ready."
            ),
            "pages_indexed": r["pages_indexed"],
            "total_chars":   r["total_chars"],
            "doc_id":        r["doc_id"],
        }
    except Exception as e:
        logger.error("Ingest error: %s", e, exc_info=True)
        raise HTTPException(500, str(e))

@app.get("/api/v1/documents")
def list_docs(): return {"documents": get_pipeline().store.list_docs()}

@app.get("/api/v1/ingest/status/{job_id}")
def ingest_status(job_id: str):
    return {"job_id": job_id, "status": "complete", "progress_pct": 100, "error": None}

@app.get("/api/v1/audit/verify/chain")
def verify_chain(): return {"chain_intact": True}

@app.get("/api/v1/audit/{audit_id}")
def get_audit(audit_id: str):
    return {"audit_id": audit_id, "timestamp": "", "user_id": "", "session_id": None,
            "compliance_status": "", "confidence_score": 0.0}

if __name__ == "__main__":
    import uvicorn
    key_ok  = bool(os.getenv("OPENAI_API_KEY", "").strip())
    api_key = os.getenv("OPENAI_API_KEY", "")
    print("\n" + "="*70)
    print("  INICAI DefenceRAG  v5.0  —  Hybrid Semantic + Keyword Search")
    print(f"  API  : http://localhost:8000")
    print(f"  Docs : http://localhost:8000/docs")
    print(f"  LLM  : {os.getenv('OPENAI_MODEL','gpt-4o-mini')} @ {os.getenv('OPENAI_BASE_URL','https://api.openai.com/v1')}")
    print(f"  Key  : {'✅ loaded (...'+api_key[-6:]+')' if key_ok else '❌ NOT SET — add OPENAI_API_KEY to .env'}")
    print("  Search: FAISS semantic (60%) + BM25 keyword (40%) → RRF fusion")
    print("  PDF  : pdfplumber tables + pymupdf text blocks")
    print("="*70 + "\n")
    uvicorn.run("dev_server:app", host="0.0.0.0", port=8000, reload=True)