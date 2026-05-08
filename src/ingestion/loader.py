# """
# src/ingestion/loader.py
# DocumentLoader: detects file type and extracts text + page structure.
# Returns a list of RawPage objects ready for the chunker.
# """
# from __future__ import annotations

# import hashlib
# import io
# import os
# from dataclasses import dataclass, field
# from pathlib import Path
# from typing import Optional

# from loguru import logger

# from config.settings import get_settings

# settings = get_settings()


# # ── Data classes ──────────────────────────────────────────────────────────────

# @dataclass
# class RawPage:
#     page_number: int
#     text: str
#     ocr_used: bool = False
#     ocr_confidence: float = 1.0


# @dataclass
# class LoadedDocument:
#     filename: str
#     file_path: str
#     content_hash: str
#     mime_type: str
#     pages: list[RawPage] = field(default_factory=list)
#     total_pages: int = 0
#     ocr_used: bool = False
#     load_error: Optional[str] = None

#     @property
#     def full_text(self) -> str:
#         return "\n\n".join(p.text for p in self.pages if p.text.strip())


# # ── Loader ────────────────────────────────────────────────────────────────────

# class DocumentLoader:
#     """
#     Routes documents to the correct parser based on MIME type.
#     Falls back to OCR for scanned PDFs where text layer is absent.
#     """

#     def __init__(self):
#         self._magic_available = self._check_magic()

#     def _check_magic(self) -> bool:
#         try:
#             import magic  # noqa
#             return True
#         except ImportError:
#             logger.warning("python-magic not available; using extension-based detection")
#             return False

#     def detect_mime_type(self, file_path: Path) -> str:
#         """Detect MIME type. Falls back to extension if python-magic unavailable."""
#         if self._magic_available:
#             try:
#                 import magic
#                 return magic.from_file(str(file_path), mime=True)
#             except Exception:
#                 pass
#         ext = file_path.suffix.lower()
#         ext_map = {
#             ".pdf": "application/pdf",
#             ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             ".doc": "application/msword",
#             ".tiff": "image/tiff",
#             ".tif": "image/tiff",
#             ".png": "image/png",
#             ".jpg": "image/jpeg",
#             ".jpeg": "image/jpeg",
#             ".txt": "text/plain",
#         }
#         return ext_map.get(ext, "application/octet-stream")

#     def compute_hash(self, file_path: Path) -> str:
#         """Compute SHA-256 of file contents for deduplication."""
#         h = hashlib.sha256()
#         with open(file_path, "rb") as f:
#             for chunk in iter(lambda: f.read(65536), b""):
#                 h.update(chunk)
#         return h.hexdigest()

#     def load(self, file_path: Path) -> LoadedDocument:
#         """Main entry point. Detects type and dispatches to correct parser."""
#         file_path = Path(file_path)
#         if not file_path.exists():
#             raise FileNotFoundError(f"File not found: {file_path}")

#         content_hash = self.compute_hash(file_path)
#         mime_type = self.detect_mime_type(file_path)

#         logger.info(f"Loading {file_path.name} | mime={mime_type} | hash={content_hash[:8]}…")

#         doc = LoadedDocument(
#             filename=file_path.name,
#             file_path=str(file_path),
#             content_hash=content_hash,
#             mime_type=mime_type,
#         )

#         try:
#             if mime_type == "application/pdf":
#                 self._load_pdf(file_path, doc)
#             elif "wordprocessingml" in mime_type or mime_type == "application/msword":
#                 self._load_docx(file_path, doc)
#             elif mime_type in ("image/tiff", "image/png", "image/jpeg"):
#                 self._load_image_ocr(file_path, doc)
#             elif mime_type == "text/plain":
#                 self._load_text(file_path, doc)
#             else:
#                 doc.load_error = f"Unsupported MIME type: {mime_type}"
#                 logger.warning(doc.load_error)
#         except Exception as e:
#             doc.load_error = str(e)
#             logger.exception(f"Failed to load {file_path.name}: {e}")

#         doc.total_pages = len(doc.pages)
#         logger.info(f"Loaded {doc.filename}: {doc.total_pages} pages, ocr={doc.ocr_used}")
#         return doc

#     # ── PDF ──────────────────────────────────────────────────────────────────

#     def _load_pdf(self, file_path: Path, doc: LoadedDocument) -> None:
#         """
#         Try PyMuPDF first for digital PDFs.
#         If text yield < 50 chars/page on average → assume scanned → OCR.
#         """
#         try:
#             import fitz  # PyMuPDF
#         except ImportError:
#             logger.warning("PyMuPDF not installed; falling back to OCR")
#             self._load_pdf_ocr(file_path, doc)
#             return

#         pdf = fitz.open(str(file_path))
#         pages: list[RawPage] = []
#         total_text_len = 0

#         for page_num in range(len(pdf)):
#             page = pdf[page_num]
#             text = page.get_text("text").strip()
#             total_text_len += len(text)
#             pages.append(RawPage(page_number=page_num + 1, text=text))

#         pdf.close()
#         avg_chars = total_text_len / max(len(pages), 1)

#         if avg_chars < 50:
#             logger.info(f"PDF avg chars/page={avg_chars:.1f} — treating as scanned, running OCR")
#             self._load_pdf_ocr(file_path, doc)
#         else:
#             # Try pdfplumber for tables (supplement PyMuPDF)
#             try:
#                 self._supplement_tables_pdfplumber(file_path, pages)
#             except Exception as e:
#                 logger.debug(f"pdfplumber table extraction skipped: {e}")
#             doc.pages = pages

#     def _supplement_tables_pdfplumber(self, file_path: Path, pages: list[RawPage]) -> None:
#         """Extract tables via pdfplumber and append to existing pages."""
#         try:
#             import pdfplumber
#         except ImportError:
#             return

#         with pdfplumber.open(str(file_path)) as pdf:
#             for i, page in enumerate(pdf.pages):
#                 if i >= len(pages):
#                     break
#                 tables = page.extract_tables()
#                 if tables:
#                     table_texts = []
#                     for table in tables:
#                         rows = [" | ".join(str(c) if c else "" for c in row) for row in table if row]
#                         table_texts.append("\n".join(rows))
#                     if table_texts:
#                         pages[i].text += "\n\n[TABLE]\n" + "\n\n".join(table_texts) + "\n[/TABLE]"

#     def _load_pdf_ocr(self, file_path: Path, doc: LoadedDocument) -> None:
#         """OCR a scanned PDF using pdf2image + pytesseract."""
#         try:
#             from pdf2image import convert_from_path
#             import pytesseract
#         except ImportError:
#             doc.load_error = "pdf2image or pytesseract not installed for OCR"
#             return

#         try:
#             images = convert_from_path(str(file_path), dpi=300)
#         except Exception as e:
#             doc.load_error = f"pdf2image failed: {e}"
#             return

#         pages = []
#         for i, img in enumerate(images):
#             data = pytesseract.image_to_data(
#                 img,
#                 lang="eng+hin",
#                 output_type=pytesseract.Output.DICT,
#                 config="--oem 3 --psm 3",
#             )
#             # Compute confidence (ignore -1 confidence entries)
#             confs = [c for c in data["conf"] if c != -1]
#             avg_conf = sum(confs) / len(confs) / 100 if confs else 0.0
#             text = pytesseract.image_to_string(img, lang="eng+hin").strip()
#             pages.append(RawPage(
#                 page_number=i + 1,
#                 text=text,
#                 ocr_used=True,
#                 ocr_confidence=avg_conf,
#             ))

#         doc.pages = pages
#         doc.ocr_used = True

#     # ── DOCX ─────────────────────────────────────────────────────────────────

#     def _load_docx(self, file_path: Path, doc: LoadedDocument) -> None:
#         try:
#             from docx import Document as DocxDocument
#         except ImportError:
#             doc.load_error = "python-docx not installed"
#             return

#         d = DocxDocument(str(file_path))
#         paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]

#         # Extract tables
#         table_texts = []
#         for table in d.tables:
#             rows = []
#             for row in table.rows:
#                 cells = " | ".join(c.text.strip() for c in row.cells)
#                 rows.append(cells)
#             table_texts.append("\n".join(rows))

#         full_text = "\n\n".join(paragraphs)
#         if table_texts:
#             full_text += "\n\n[TABLE]\n" + "\n\n".join(table_texts) + "\n[/TABLE]"

#         # DOCX has no concept of pages — treat as single page
#         doc.pages = [RawPage(page_number=1, text=full_text)]

#     # ── Image OCR ─────────────────────────────────────────────────────────────

#     def _load_image_ocr(self, file_path: Path, doc: LoadedDocument) -> None:
#         try:
#             from PIL import Image
#             import pytesseract
#         except ImportError:
#             doc.load_error = "Pillow or pytesseract not installed"
#             return

#         img = Image.open(str(file_path))
#         data = pytesseract.image_to_data(
#             img, lang="eng+hin",
#             output_type=pytesseract.Output.DICT,
#         )
#         confs = [c for c in data["conf"] if c != -1]
#         avg_conf = sum(confs) / len(confs) / 100 if confs else 0.0
#         text = pytesseract.image_to_string(img, lang="eng+hin").strip()
#         doc.pages = [RawPage(page_number=1, text=text, ocr_used=True, ocr_confidence=avg_conf)]
#         doc.ocr_used = True

#     # ── Plain text ────────────────────────────────────────────────────────────

#     def _load_text(self, file_path: Path, doc: LoadedDocument) -> None:
#         text = file_path.read_text(encoding="utf-8", errors="replace")
#         doc.pages = [RawPage(page_number=1, text=text)]
"""
loader.py — Document loader for PDF and DOCX files.

Returns plain text extracted from the document.
Falls back gracefully if optional libraries are missing.
"""

from __future__ import annotations
import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    file_path: str
    filename: str
    text: str
    page_count: int = 0
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """
    Loads PDF and DOCX files into plain text.
    Tries multiple backends in order of preference.
    """

    SUPPORTED = {".pdf", ".docx", ".doc", ".txt"}

    def load(self, file_path: str) -> LoadedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED:
            raise ValueError(f"Unsupported file type: {suffix}. Supported: {self.SUPPORTED}")

        logger.info("Loading document: %s", path.name)

        if suffix == ".pdf":
            return self._load_pdf(path)
        elif suffix in (".docx", ".doc"):
            return self._load_docx(path)
        elif suffix == ".txt":
            return self._load_txt(path)
        else:
            raise ValueError(f"No loader for: {suffix}")

    # ------------------------------------------------------------------

    def _load_pdf(self, path: Path) -> LoadedDocument:
        text = ""
        page_count = 0

        # Try PyMuPDF (fitz) first — fastest and best quality
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            pages = []
            for page in doc:
                pages.append(page.get_text())
            text = "\n\n".join(pages)
            page_count = len(doc)
            doc.close()
            logger.info("PDF loaded via PyMuPDF: %d pages", page_count)
        except ImportError:
            # Fallback: pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    pages = []
                    for page in pdf.pages:
                        t = page.extract_text() or ""
                        pages.append(t)
                    text = "\n\n".join(pages)
                    page_count = len(pdf.pages)
                logger.info("PDF loaded via pdfplumber: %d pages", page_count)
            except ImportError:
                # Last fallback: pypdf
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(str(path))
                    pages = []
                    for page in reader.pages:
                        pages.append(page.extract_text() or "")
                    text = "\n\n".join(pages)
                    page_count = len(reader.pages)
                    logger.info("PDF loaded via pypdf: %d pages", page_count)
                except ImportError:
                    raise ImportError(
                        "No PDF library found. Install one of: pymupdf, pdfplumber, pypdf\n"
                        "  pip install pymupdf"
                    )

        text = self._clean_text(text)
        return LoadedDocument(
            file_path=str(path),
            filename=path.name,
            text=text,
            page_count=page_count,
            metadata={"source": path.name, "type": "pdf"},
        )

    def _load_docx(self, path: Path) -> LoadedDocument:
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            logger.info("DOCX loaded: %d paragraphs", len(paragraphs))
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        text = self._clean_text(text)
        return LoadedDocument(
            file_path=str(path),
            filename=path.name,
            text=text,
            page_count=0,
            metadata={"source": path.name, "type": "docx"},
        )

    def _load_txt(self, path: Path) -> LoadedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        text = self._clean_text(text)
        return LoadedDocument(
            file_path=str(path),
            filename=path.name,
            text=text,
            page_count=1,
            metadata={"source": path.name, "type": "txt"},
        )

    @staticmethod
    def _clean_text(text: str) -> str:
        import re
        # Collapse excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        return text.strip()