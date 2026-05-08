"""
Document loading helpers for the lightweight RAG pipeline.

The loader favors PyMuPDF for PDFs because it is fast and already present in
the project dependencies. DOCX and TXT are handled locally. The returned shape
is intentionally small so ingestion stays predictable.
"""

from __future__ import annotations

import hashlib
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    file_path: str
    filename: str
    text: str
    page_count: int = 0
    content_hash: str = ""
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """Load supported document formats into clean text."""

    SUPPORTED = {".pdf", ".docx", ".doc", ".txt"}

    def load(self, file_path: str) -> LoadedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED:
            raise ValueError(f"Unsupported file type: {suffix}. Supported: {sorted(self.SUPPORTED)}")

        logger.info("Loading document: %s", path.name)
        content_hash = self.compute_hash(path)

        if suffix == ".pdf":
            return self._load_pdf(path, content_hash)
        if suffix in (".docx", ".doc"):
            return self._load_docx(path, content_hash)
        return self._load_txt(path, content_hash)

    @staticmethod
    def compute_hash(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as fh:
            for chunk in iter(lambda: fh.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    def _load_pdf(self, path: Path, content_hash: str) -> LoadedDocument:
        try:
            import fitz
        except ImportError as exc:
            raise ImportError("PyMuPDF is required for PDF ingestion. Install with: pip install pymupdf") from exc

        pages: list[str] = []
        with fitz.open(str(path)) as doc:
            for page_number, page in enumerate(doc, start=1):
                page_text = page.get_text("text") or ""
                page_text = self._clean_text(page_text)
                if page_text:
                    pages.append(f"[PAGE {page_number}]\n{page_text}")
            page_count = len(doc)

        text = "\n\n".join(pages)
        logger.info("PDF loaded via PyMuPDF: %d pages, %d chars", page_count, len(text))
        return LoadedDocument(
            file_path=str(path),
            filename=path.name,
            text=text,
            page_count=page_count,
            content_hash=content_hash,
            metadata={"source": path.name, "type": "pdf", "content_hash": content_hash},
        )

    def _load_docx(self, path: Path, content_hash: str) -> LoadedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("python-docx is required for DOCX ingestion. Install with: pip install python-docx") from exc

        doc = Document(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append("[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")

        text = self._clean_text("\n\n".join(parts))
        logger.info("DOCX loaded: %d text blocks, %d chars", len(parts), len(text))
        return LoadedDocument(
            file_path=str(path),
            filename=path.name,
            text=text,
            page_count=0,
            content_hash=content_hash,
            metadata={"source": path.name, "type": "docx", "content_hash": content_hash},
        )

    def _load_txt(self, path: Path, content_hash: str) -> LoadedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")
        text = self._clean_text(text)
        return LoadedDocument(
            file_path=str(path),
            filename=path.name,
            text=text,
            page_count=1,
            content_hash=content_hash,
            metadata={"source": path.name, "type": "txt", "content_hash": content_hash},
        )

    @staticmethod
    def _clean_text(text: str) -> str:
        text = text.replace("\x00", " ")
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
